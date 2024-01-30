"""Models for the Reports."""
import io
import os
from os.path import splitext
from uuid import uuid4
from zipfile import BadZipFile, ZipFile

from django.contrib.auth.models import User
from django.core.exceptions import ValidationError
from django.core.files.base import ContentFile
from django.db import models
from django.db.models.signals import post_save

# for converting docx to pdf
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate


def convert_docx_to_pdf(docx_path):
    """Convert docx to pdf."""
    # using python-docx to read the DOCX file
    doc = Document(docx_path)

    # setting the page size and margins
    left_margin = 1 * inch
    right_margin = 1 * inch

    # creating a PDF using reportlab
    pdf_path = "media/reports/new_filename.pdf"
    pdf = SimpleDocTemplate(
        pdf_path, pagesize=letter, leftMargin=left_margin, rightMargin=right_margin
    )
    # using python-docx to read the DOCX file
    doc = Document(docx_path)

    # defining some styles
    styles = getSampleStyleSheet()
    # adding custome styles if they do not yet exist
    if "Heading1" not in styles:
        styles.add(ParagraphStyle(name="Heading1", fontSize=18, leading=22))
    if "BodyText" not in styles:
        styles.add(ParagraphStyle(name="BodyText", fontSize=12, leading=16))

    # creating list of content that goes into the pdf
    pdf_content = []
    for paragraph in doc.paragraphs:
        style = None
        if paragraph.style.name == "Heading 1":
            style = "Heading1"
        else:
            style = "BodyText"
        pdf_content.append(Paragraph(paragraph.text, styles[style]))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                pdf_content.append(Paragraph(cell.text, styles["BodyText"]))

    # building the pdf
    pdf.build(pdf_content)
    # reading the generated pdf to return its data
    with open(pdf_path, "rb") as f:
        pdf_data = f.read()
    return pdf_data


def convert_and_save(instance):
    """Convert and save a PDF instance."""
    uploaded_file = instance.file
    initial_bytes = uploaded_file.read(20)
    uploaded_file.seek(0)
    # checking if file is docx
    is_docx = initial_bytes.startswith(b"PK\x03\x04")
    if is_docx:
        docx_path = uploaded_file.path
        pdf_data = convert_docx_to_pdf(docx_path)
        # deleting the original docx file
        os.remove(docx_path)
        # generating a PDF file name based on the instance name or UUID
        pdf_name = f"{instance.name}.pdf" if instance.name else f"{uuid4()}.pdf"
        # saving the new pdf content directly into the database
        instance.file.save(
            pdf_name, ContentFile(pdf_data)
        )  # update current instance of pdf file
        instance.save()
        # removing the intermediate PDF file if it exists
        intermediate_pdf_path = "media/reports/new_filename.pdf"
        if os.path.exists(intermediate_pdf_path):
            os.remove(intermediate_pdf_path)


def report_post_save(sender, instance, created, **kwargs):
    """Handle the signal for post-save actions."""
    if created:
        convert_and_save(instance)  # convert and save file


def file_upload_path(instance, filename):
    """Calculate a randomised file upload path for a Report."""
    # file will be uploaded to MEDIA_ROOT/reports/<filename>
    return "reports/" + str(uuid4()) + splitext(filename)[-1]


def validate_file(file):
    """Validate whether the file is actually a report."""
    MAX_FILE_SIZE = 1024 * 1024 * 5  # 5 MB
    if file is None:
        raise ValidationError("Missing file.")
    # Check file size
    if file.size > MAX_FILE_SIZE:
        raise ValidationError("File size must be no more than 5 MB.")
    # Read initial bytes
    initial_bytes = file.read(20)  # Read the first 20 bytes
    file.seek(0)
    # Check for PDF
    is_pdf = initial_bytes.startswith(b"%PDF")  # Removed the newline check
    # Check for DOCX
    is_docx = False
    if initial_bytes.startswith(b"PK\x03\x04"):
        try:
            with ZipFile(io.BytesIO(file.read())) as myzip:
                is_docx = "[Content_Types].xml" in myzip.namelist()
        except BadZipFile:
            is_docx = False
        finally:
            file.seek(0)  # Reset file to its original state
    if not (is_pdf or is_docx):
        raise ValidationError(
            "Unsupported file type. Only PDF and DOCX files are allowed."
        )


class Report(models.Model):
    """Model (or table (in MySQL equiv)) for the Reports."""

    # deleting user automatically deletes associated reports
    user = models.ForeignKey(User, null=True, on_delete=models.CASCADE)
    name = models.CharField(max_length=64, unique=True)
    # variable which holds the path of the file, i.e., reports/document_name.filetype
    # validator called first
    file = models.FileField(upload_to=file_upload_path, validators=[validate_file])
    # variable holds the time of upload (aiding in version control)
    last_updated = models.DateTimeField(auto_now=True)

    def __str__(self):
        """Return a string representation for a Report."""
        return self.name


post_save.connect(
    report_post_save, sender=Report
)  # connects post_save signal to the report model
